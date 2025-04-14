import re
# 在脚本开头添加
# import matplotlib.font_manager as fm
#fm._rebuild()  # 重建字体缓存
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patheffects as path_effects
from matplotlib.ticker import MultipleLocator, FormatStrFormatter, FixedLocator
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from collections import defaultdict
import os
import sys
import pandas as pd
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 添加这个导入
import os
from bs4 import BeautifulSoup
def create_html_table(models_dict, output_file='model_comparison_table.html'):
    """
    创建一个HTML表格并保存到文件
    
    Args:
        models_dict: 字典，键是模型名称，值是包含模型信息的字典列表
        output_file: 输出HTML文件路径
    """
    # 创建表格数据
    table_data = []
    model_rowspans = {}  # 存储每个模型名称需要合并的行数
    
    for model_name, model_variants in models_dict.items():
        # 记录该模型的变体数量（用于行合并）
        if model_name == "Unknown":
            continue
        model_rowspans[model_name] = len(model_variants)
        
        for i, variant in enumerate(model_variants):
            row = []
            
            # 只在第一行添加模型名称
            if i == 0:
                row.append(model_name)
            else:
                row.append("")  # 空单元格用于后续合并
            
            # 添加其他列 - 按照新的顺序
            row.append(variant.get('quantization', ''))  # 量化版本
            row.append(variant.get('model_type', ''))   # 模型类型
            row.append(variant.get('bpw', ''))          # BPW
            row.append(variant.get('model_size_gb', '')) # 模型大小
            
            table_data.append(row)
    
    # 创建DataFrame
    df = pd.DataFrame(table_data, columns=['Model Name', 'Quantization', 'Model Type', 'BPW', 'Model Size (GB)'])
    
    # 生成基本HTML
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Model Comparison</title>
        <style>
            table {
                border-collapse: collapse;
                width: 100%;
            }
            th, td {
                border: 1px solid #dddddd;
                text-align: left;
                padding: 8px;
            }
            th {
                background-color: #f2f2f2;
            }
            tr:nth-child(even) {
                background-color: #f9f9f9;
            }
        </style>
    </head>
    <body>
        <h2>Model Comparison Table</h2>
        <table>
            <tr>
                <th>Model Name</th>
                <th>Quantization</th>
                <th>Model Type</th>
                <th>BPW</th>
                <th>Model Size (GB)</th>
            </tr>
    """
    
    # 添加表格数据行，处理行合并
    current_model = None
    for i, row in enumerate(table_data):
        html_content += "<tr>"
        
        for j, cell in enumerate(row):
            # 处理第一列（模型名称）的单元格合并
            if j == 0:
                if cell and cell != current_model:
                    current_model = cell
                    rowspan = model_rowspans[cell]
                    html_content += f'<td rowspan="{rowspan}">{cell}</td>'
                # 如果是空单元格（已被上面的单元格合并），则跳过
                elif not cell or cell == "":
                    continue
                else:
                    html_content += f'<td>{cell}</td>'
            else:
                html_content += f'<td>{cell}</td>'
                
        html_content += "</tr>\n"
    
    html_content += """
        </table>
    </body>
    </html>
    """
    
    # 保存到文件
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"HTML表格已保存到 {os.path.abspath(output_file)}")
    return html_content

def html_table_to_word(html_content, output_file='model_comparison_from_html.docx'):
    """
    将HTML表格转换为Word表格
    
    Args:
        html_content: HTML表格内容或HTML文件路径
        output_file: 输出Word文件路径
    """
    # 如果输入是文件路径，读取文件内容
    if os.path.isfile(html_content):
        with open(html_content, 'r', encoding='utf-8') as f:
            html_content = f.read()
    
    # 解析HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 查找表格
    html_table = soup.find('table')
    if not html_table:
        print("未找到HTML表格")
        return
    
    # 查找表头
    headers = []
    header_row = html_table.find('tr')
    if header_row:
        headers = [th.get_text().strip() for th in header_row.find_all(['th', 'td'])]
    
    # 创建Word文档
    doc = Document()
    doc.add_heading('Model Comparison Table', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 查找所有数据行
    rows = html_table.find_all('tr')
    if len(rows) <= 1:  # 只有表头或没有行
        print("表格没有数据行")
        return
    
    # 提取行和列数
    num_rows = len(rows)
    num_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
    
    # 创建Word表格
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # 处理表头
    if headers:
        for i, header in enumerate(headers):
            if i < num_cols:  # 确保不超出列数
                cell = table.cell(0, i)
                cell.text = header
                # 设置表头格式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # 处理数据行和合并单元格
    rowspan_tracking = {}  # 跟踪需要跳过的单元格
    
    for i, row in enumerate(rows):
        if i == 0 and headers:  # 跳过表头行，因为已经处理过了
            continue
            
        cells = row.find_all(['th', 'td'])
        word_row_idx = i if not headers else i - 1
        
        col_idx = 0
        for j, cell in enumerate(cells):
            # 检查是否需要跳过此单元格（被上方单元格合并）
            while (word_row_idx, col_idx) in rowspan_tracking:
                col_idx += 1
            
            # 获取rowspan和colspan属性
            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))
            
            # 如果有rowspan，标记需要跳过的单元格
            if rowspan > 1:
                for rs in range(1, rowspan):
                    for cs in range(colspan):
                        rowspan_tracking[(word_row_idx + rs, col_idx + cs)] = True
            
            # 设置单元格文本
            if word_row_idx < table.rows.__len__() and col_idx < num_cols:
                word_cell = table.cell(word_row_idx, col_idx)
                word_cell.text = cell.get_text().strip()
            
            # 处理colspan（合并列）
            if colspan > 1 and word_row_idx < table.rows.__len__():
                for cs in range(1, colspan):
                    if col_idx + cs < num_cols:
                        word_cell.merge(table.cell(word_row_idx, col_idx + cs))
            
            col_idx += colspan
    
    # 保存文档
    doc.save(output_file)
    print(f"Word表格已保存到 {os.path.abspath(output_file)}")
    return output_file

def create_html_and_word_tables(models_dict):
    """创建HTML和Word格式的表格"""
    # 生成HTML表格
    html_file = 'model_comparison_table.html'
    html_content = create_html_table(models_dict, html_file)
    
    # 将HTML转换为Word
    word_file = 'model_comparison_table.docx'
    html_table_to_word(html_file, word_file)
    
    return html_file, word_file
# 使用示例
# html_file, word_file = create_html_and_word_tables(models_dict)
def create_word_friendly_html_table(models_dict):
    """
    创建可直接复制到Word的HTML表格
    
    Args:
        models_dict: 字典，键是模型名称，值是包含模型信息的字典列表
    
    Returns:
        HTML表格字符串，可直接复制到Word
    """
    # 生成表格HTML
    html_content = """
    <table border="1" cellspacing="0" cellpadding="5" style="border-collapse: collapse;">
        <tr style="background-color: #f2f2f2; font-weight: bold;">
            <td>Model Name</td>
            <td>Quantization</td>
            <td>Bit Weight</td>
            <td>Model Size (GB)</td>
        </tr>
    """
    
    # 跟踪需要合并的单元格
    current_model = None
    model_rowspans = {}
    
    # 计算每个模型的行数
    for model_name, variants in models_dict.items():
        model_rowspans[model_name] = len(variants)
    
    # 添加数据行
    for model_name, variants in models_dict.items():
        for i, variant in enumerate(variants):
            html_content += "<tr>"
            
            # 处理模型名称列（第一列）
            if i == 0:  # 只在第一个变体行添加模型名称
                rowspan = model_rowspans[model_name]
                if rowspan > 1:
                    html_content += f'<td rowspan="{rowspan}" valign="center">{model_name}</td>'
                else:
                    html_content += f'<td>{model_name}</td>'
            # 对于同一模型的其他变体，不添加模型名称单元格（已被合并）
            
            # 添加其他列
            html_content += f'<td>{variant.get("quantization", "")}</td>'
            html_content += f'<td>{variant.get("model_ftype", "")}</td>'
            html_content += f'<td>{variant.get("model_size_gb", "")}</td>'
            
            html_content += "</tr>\n"
    
    html_content += "</table>"
    return html_content
def generate_figure(models_data):
    df = pd.DataFrame(models_data)

    # 计算token生成时间
    df['token_generation_time'] = df['total_time_ms'] / df['total_tokens']

    # 自定义排序函数，确保F16在末尾
    def quantization_sort_key(quant_str):
        if 'F16' in quant_str or 'f16' in quant_str:
            return 100  # 确保F16在末尾
        match = re.search(r'Q(\d+)', quant_str)
        if match:
            return int(match.group(1))
        return 99  # 其他格式放在F16之前

    # 添加排序键
    df['quant_sort_key'] = df['quantization'].apply(quantization_sort_key)

    # 按量化级别排序
    df = df.sort_values('quant_sort_key')

    # 提取唯一的量化方法并排序
    quant_methods = df['quantization'].unique()
    
    # 设置中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans', 'Arial']
    plt.rcParams['axes.unicode_minus'] = False

    # 设置不同模型的标记和颜色
    base_models = df['base_model'].unique()
    markers = ['o', 's', '^', 'D', 'v', '*', 'p', 'h']
    colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f']

    # 创建模型到标记和颜色的映射
    model_markers = {model: markers[i % len(markers)] for i, model in enumerate(base_models)}
    model_colors = {model: colors[i % len(colors)] for i, model in enumerate(base_models)}

    # 创建图形
    fig_width = 12
    fig_height = 6

    # 1. 绘制困惑度(PPL)对比图
    plt.figure(figsize=(fig_width, fig_height))
    
    # 创建横轴位置
    x_positions = np.arange(len(quant_methods)) * 1.5  # 增加间距
    
    # 创建x轴标签位置的映射
    x_ticks_map = {method: pos for method, pos in zip(quant_methods, x_positions)}
    
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            plt.scatter(
                x_ticks_map[row['quantization']], 
                row['ppl'],
                marker=model_markers[model], 
                color=model_colors[model],
                s=100,  # 增大点的大小
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""  # 只为每个模型添加一次图例
            )

    plt.title('不同量化方法对困惑度(PPL)的影响', fontsize=16)
    plt.xlabel('量化方法', fontsize=14)
    plt.ylabel('困惑度 (PPL)', fontsize=14)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.xticks(x_positions, quant_methods, rotation=45)
    
    # 去除重复的图例
    handles, labels = plt.gca().get_legend_handles_labels()
    unique = [(h, l) for i, (h, l) in enumerate(zip(handles, labels)) if l not in labels[:i]]
    plt.legend(*zip(*unique), title='基础模型', fontsize=12)
    
    plt.tight_layout()
    plt.savefig('困惑度对比图.png', dpi=300, bbox_inches='tight')
    plt.close()

    # 2. 绘制加载时间对比图
    plt.figure(figsize=(fig_width, fig_height))
    
    # 获取加载时间数据范围
    load_times_sec = df['load_time_ms'].values / 1000
    max_load_time = max(load_times_sec)
    
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            plt.scatter(
                x_ticks_map[row['quantization']], 
                row['load_time_ms'] / 1000,
                marker=model_markers[model], 
                color=model_colors[model],
                s=100,  # 增大点的大小
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""
            )

    plt.title('不同量化方法对模型加载时间的影响', fontsize=16)
    plt.xlabel('量化方法', fontsize=14)
    plt.ylabel('加载时间 (秒)', fontsize=14)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.xticks(x_positions, quant_methods, rotation=45)
    
    # 调整Y轴刻度，增加0-100秒区间的区分度
    if max_load_time > 100:
        # 使用非线性刻度，增加底部区间的区分度
        plt.yscale('symlog', linthresh=100)  # linthresh定义线性区域的阈值
        
        # 手动设置刻度，增加0-100区间的刻度密度
        y_ticks = list(range(0, 101, 20)) + list(range(200, int(max_load_time) + 100, 100))
        plt.yticks(y_ticks)
    else:
        # 如果所有数据都在100秒以内，使用线性刻度
        plt.ylim(0, max(100, max_load_time * 1.1))
        
        # 增加刻度密度
        plt.gca().yaxis.set_major_locator(MultipleLocator(10))  # 每10秒一个主刻度
        plt.gca().yaxis.set_minor_locator(MultipleLocator(5))   # 每5秒一个次刻度
    
    # 去除重复的图例
    handles, labels = plt.gca().get_legend_handles_labels()
    unique = [(h, l) for i, (h, l) in enumerate(zip(handles, labels)) if l not in labels[:i]]
    plt.legend(*zip(*unique), title='基础模型', fontsize=12)
    
    plt.tight_layout()
    plt.savefig('加载时间对比图.png', dpi=300, bbox_inches='tight')
    plt.close()

    # 3. 绘制Token生成时间对比图
    plt.figure(figsize=(fig_width, fig_height))
    
    # 获取token生成时间数据范围
    token_times = df['token_generation_time'].values
    max_token_time = max(token_times)
    
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            plt.scatter(
                x_ticks_map[row['quantization']], 
                row['token_generation_time'],
                marker=model_markers[model], 
                color=model_colors[model],
                s=100,  # 增大点的大小
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""
            )

    plt.title('不同量化方法对Token生成时间的影响', fontsize=16)
    plt.xlabel('量化方法', fontsize=14)
    plt.ylabel('每Token生成时间 (毫秒)', fontsize=14)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.xticks(x_positions, quant_methods, rotation=45)
    
    # 调整Y轴刻度，增加0-200毫秒区间的区分度
    if max_token_time > 200:
        # 使用非线性刻度，增加底部区间的区分度
        plt.yscale('symlog', linthresh=200)  # linthresh定义线性区域的阈值
        
        # 手动设置刻度，增加0-200区间的刻度密度
        y_ticks = list(range(0, 201, 25)) + list(range(300, int(max_token_time) + 100, 100))
        plt.yticks(y_ticks)
    else:
        # 如果所有数据都在200毫秒以内，使用线性刻度
        plt.ylim(0, max(200, max_token_time * 1.1))
        
        # 增加刻度密度
        plt.gca().yaxis.set_major_locator(MultipleLocator(25))  # 每25毫秒一个主刻度
        plt.gca().yaxis.set_minor_locator(MultipleLocator(5))   # 每5毫秒一个次刻度
    
    # 去除重复的图例
    handles, labels = plt.gca().get_legend_handles_labels()
    unique = [(h, l) for i, (h, l) in enumerate(zip(handles, labels)) if l not in labels[:i]]
    plt.legend(*zip(*unique), title='基础模型', fontsize=12)
    
    plt.tight_layout()
    plt.savefig('Token生成时间对比图.png', dpi=300, bbox_inches='tight')
    plt.close()

    # 4. 绘制综合对比图（三个指标在一起）
    fig, axes = plt.subplots(1, 3, figsize=(18, 6))

    # PPL对比
    ax = axes[0]
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            ax.scatter(
                x_ticks_map[row['quantization']], 
                row['ppl'],
                marker=model_markers[model], 
                color=model_colors[model],
                s=80,  # 稍微小一点的点
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""
            )
        
    ax.set_title('困惑度(PPL)对比', fontsize=14)
    ax.set_xlabel('量化方法', fontsize=12)
    ax.set_ylabel('困惑度 (PPL)', fontsize=12)
    ax.grid(True, linestyle='--', alpha=0.7)
    ax.set_xticks(x_positions)
    ax.set_xticklabels(quant_methods, rotation=45)

    # 加载时间对比
    ax = axes[1]
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            ax.scatter(
                x_ticks_map[row['quantization']], 
                row['load_time_ms'] / 1000,
                marker=model_markers[model], 
                color=model_colors[model],
                s=80,
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""
            )
        
    ax.set_title('模型加载时间对比', fontsize=14)
    ax.set_xlabel('量化方法', fontsize=12)
    ax.set_ylabel('加载时间 (秒)', fontsize=12)
    ax.grid(True, linestyle='--', alpha=0.7)
    ax.set_xticks(x_positions)
    ax.set_xticklabels(quant_methods, rotation=45)
    
    # 调整Y轴刻度，增加0-100秒区间的区分度
    if max_load_time > 100:
        # 使用非线性刻度，增加底部区间的区分度
        ax.set_yscale('symlog', linthresh=100)  # linthresh定义线性区域的阈值
        
        # 手动设置刻度，增加0-100区间的刻度密度
        y_ticks = list(range(0, 101, 20)) + list(range(200, int(max_load_time) + 100, 100))
        ax.set_yticks(y_ticks)
    else:
        # 如果所有数据都在100秒以内，使用线性刻度
        ax.set_ylim(0, max(100, max_load_time * 1.1))
        
        # 增加刻度密度
        ax.yaxis.set_major_locator(MultipleLocator(10))  # 每10秒一个主刻度
        ax.yaxis.set_minor_locator(MultipleLocator(5))   # 每5秒一个次刻度

    # Token生成时间对比
    ax = axes[2]
    for model in base_models:
        model_data = df[df['base_model'] == model]
        # 仅绘制散点，不连线
        for _, row in model_data.iterrows():
            ax.scatter(
                x_ticks_map[row['quantization']], 
                row['token_generation_time'],
                marker=model_markers[model], 
                color=model_colors[model],
                s=80,
                label=model if row['quantization'] == model_data['quantization'].iloc[0] else ""
            )
        
    ax.set_title('Token生成时间对比', fontsize=14)
    ax.set_xlabel('量化方法', fontsize=12)
    ax.set_ylabel('每Token生成时间 (毫秒)', fontsize=12)
    ax.grid(True, linestyle='--', alpha=0.7)
    ax.set_xticks(x_positions)
    ax.set_xticklabels(quant_methods, rotation=45)
    
    # 调整Y轴刻度，增加0-200毫秒区间的区分度
    if max_token_time > 200:
        # 使用非线性刻度，增加底部区间的区分度
        ax.set_yscale('symlog', linthresh=200)  # linthresh定义线性区域的阈值
        
        # 手动设置刻度，增加0-200区间的刻度密度
        y_ticks = list(range(0, 201, 25)) + list(range(300, int(max_token_time) + 100, 100))
        ax.set_yticks(y_ticks)
    else:
        # 如果所有数据都在200毫秒以内，使用线性刻度
        ax.set_ylim(0, max(200, max_token_time * 1.1))
        
        # 增加刻度密度
        ax.yaxis.set_major_locator(MultipleLocator(25))  # 每25毫秒一个主刻度
        ax.yaxis.set_minor_locator(MultipleLocator(5))   # 每5毫秒一个次刻度

    # 只在最后一个子图上添加图例
    handles, labels = [], []
    for model in base_models:
        handles.append(plt.Line2D([0], [0], marker=model_markers[model], color=model_colors[model], 
                                 linestyle='None', markersize=8, label=model))
        labels.append(model)
    
    fig.legend(handles, labels, loc='upper center', bbox_to_anchor=(0.5, 0), 
            ncol=len(base_models), title='基础模型', fontsize=12)

    plt.tight_layout()
    plt.subplots_adjust(bottom=0.2)  # 为图例留出空间
    plt.savefig('量化方法综合性能对比图.png', dpi=300, bbox_inches='tight')
    plt.show()

def parse_log_file(log_content):
    """
    解析日志内容，提取每个模型的指标
    
    Args:
        log_content: 日志文件的内容字符串
    
    Returns:
        list: 包含每个模型指标字典的列表
    """
    # 使用正则表达式分割不同模型的数据块
    model_blocks = re.split(r'\n(?=[\w\-\.]+\.gguf)', log_content)
    
    models_data = []
    
    for block in model_blocks:
        if not block.strip():
            continue
            
        model_data = {}
        
        # 逐行处理
        lines = block.strip().split('\n')
        for line in lines:
            line = line.strip()
            
            # 提取模型名称
            if not line.startswith('model ') and '.gguf' in line:
                model_name = line
                model_data['model_name'] = model_name
                
                # 从模型名称中提取基本模型名称和量化版本
                name_parts = model_name.split('-')
                if len(name_parts) >= 2:
                    # 提取基本模型名称 (不包括量化后缀)
                    model_data['base_model'] = '-'.join(name_parts[:-1])
                    
                    # 提取量化版本
                    quant_part = name_parts[-1]
                    quant_match = re.search(r'(Q\d+.*?)\.gguf', quant_part)
                    if quant_match:
                        model_data['quantization'] = quant_match.group(1)
                    else:
                        model_data['quantization'] = quant_part.split('.')[0]
            
            # 提取模型类型
            elif line.startswith('model type ='):
                model_data['model_type'] = line.split('=')[1].strip()
            
            # 提取模型格式类型
            elif line.startswith('model ftype ='):
                model_data['model_ftype'] = line.split('=')[1].strip()
            
            # 提取模型参数
            elif line.startswith('model params ='):
                params_part = line.split('=')[1].strip()
                params_match = re.match(r'([\d\.]+) (\w)', params_part)
                if params_match:
                    value = float(params_match.group(1))
                    unit = params_match.group(2)
                    if unit == 'B':
                        model_data['model_params'] = value * 1e9
                    elif unit == 'M':
                        model_data['model_params'] = value * 1e6
                    elif unit == 'K':
                        model_data['model_params'] = value * 1e3
                    else:
                        model_data['model_params'] = value
            
            # 提取模型大小和 BPW - 使用字符串处理
            elif line.startswith('model size ='):
                # 提取模型大小
                size_part = line.split('=')[1].strip()
                size_match = re.match(r'([\d\.]+) (\w+)', size_part)
                if size_match:
                    value = float(size_match.group(1))
                    unit = size_match.group(2)
                    if unit == 'GiB':
                        model_data['model_size_gb'] = value
                    elif unit == 'MiB':
                        model_data['model_size_gb'] = value / 1024.0
                    else:
                        model_data['model_size_gb'] = value
                
                # 提取 BPW - 查找括号并提取内容
                if '(' in line and ')' in line:
                    bpw_part = line.split('(')[1].split(')')[0]
                    bpw_match = re.search(r'([\d\.]+) BPW', bpw_part)
                    if bpw_match:
                        model_data['bpw'] = float(bpw_match.group(1))
            
            # 提取 PPL (Perplexity)
            elif line.startswith('Final estimate:'):
                ppl_match = re.search(r'PPL = ([\d\.]+) \+/- ([\d\.]+)', line)
                if ppl_match:
                    model_data['ppl'] = float(ppl_match.group(1))
                    model_data['ppl_error'] = float(ppl_match.group(2))
            
            # 提取加载时间
            elif line.startswith('load time ='):
                time_part = line.split('=')[1].strip()
                time_match = re.match(r'\s*([\d\.]+) ms', time_part)
                if time_match:
                    model_data['load_time_ms'] = float(time_match.group(1))
            
            # 提取 prompt 评估时间
            elif line.startswith('prompt eval time ='):
                eval_match = re.search(r'prompt eval time = \s*([\d\.]+) ms / \s*(\d+) tokens $$\s*([\d\.]+) ms per token,\s*([\d\.]+) tokens per second$$', line)
                if eval_match:
                    model_data['prompt_eval_time_ms'] = float(eval_match.group(1))
                    model_data['prompt_tokens'] = int(eval_match.group(2))
                    model_data['ms_per_token'] = float(eval_match.group(3))
                    model_data['tokens_per_second'] = float(eval_match.group(4))
            
            # 提取总时间
            elif line.startswith('total time ='):
                total_match = re.search(r'total time = \s*([\d\.]+) ms / \s*(\d+) tokens', line)
                if total_match:
                    model_data['total_time_ms'] = float(total_match.group(1))
                    model_data['total_tokens'] = int(total_match.group(2))
        
        if model_data:  # 确保模型数据不为空
            models_data.append(model_data)
    
    return models_data
def group_models_by_base_and_quant(models_data):
    """
    将模型按基础模型和量化版本分组
    """
    grouped_data = defaultdict(list)
    
    for model in models_data:
        base_model = model.get('base_model', 'Unknown')
        grouped_data[base_model].append(model)
    
    return grouped_data

def process_log_file(log_file_path):
    """
    处理日志文件并生成可视化
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(log_file_path):
            print(f"错误: 文件 '{log_file_path}' 不存在")
            return [], {}
        
        # 读取日志文件
        with open(log_file_path, 'r', encoding='utf-8') as f:
            log_content = f.read()
        
        # 解析日志内容
        models_data = parse_log_file(log_content)
        
        # 打印提取的数据
        print(f"共提取到 {len(models_data)} 个模型的信息:")
        for model in models_data:
            print(f"\n{model.get('model_name', 'Unknown')} 的指标:")
            for key, value in model.items():
                print(f"  {key}: {value}")
        
        # 转换为 DataFrame 并打印列名
        df = pd.DataFrame(models_data)
        print("\nDataFrame 的列名:", df.columns.tolist())
        
        if df.empty:
            print("警告: 未能从日志文件中提取有效数据")
            return [], {}
        
        # 将模型按基础模型分组
        grouped_models = group_models_by_base_and_quant(models_data)
        
        # 打印分组结果
        print("\n按基础模型分组结果:")
        for base_model, models in grouped_models.items():
            print(f"\n{base_model} 系列模型:")
            for model in models:
                print(f"  - {model.get('model_name', 'Unknown')} ({model.get('quantization', 'Unknown')})")
        
        return models_data, grouped_models
    
    except Exception as e:
        print(f"处理日志文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return [], {}

if __name__ == "__main__":
    # 设置 matplotlib 使用非交互式后端，避免在无 GUI 环境中出错
    plt.switch_backend('agg')
    
    if len(sys.argv) > 1:
        log_file_path = sys.argv[1]
    else:
        log_file_path = "metrics/model_perplexity_metric.log"  # 默认日志文件路径
    
    print(f"正在处理日志文件: {log_file_path}")
    models_data, grouped_models = process_log_file(log_file_path)
    # res = create_html_and_word_tables(grouped_models)
    # pass
    # 绘制模型大小与PPL关系图
    if models_data:
        generate_figure(models_data)

    else:
        print("未能从日志文件中提取有效数据，无法生成分析结果")